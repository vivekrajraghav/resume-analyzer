import os
import time
import pypdf
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pathlib import Path
from pypdf import PdfReader
from docx import Document

# vairbale from .env
load_dotenv()

# Getting API key from .env
my_api_key=os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError("API key is missing")

# Creating Client and selecting model
client=Groq(api_key=my_api_key)
model="openai/gpt-oss-120b"

job_description="""
About the job
We’re looking for a Machine Learning Intern who’s excited about building models that solve real Trust & Safety problems at scale. 🚀



As part of the FRND ML team, you’ll train and ship models that help identify fake profiles, unsafe interactions, abusive speech, harassment, spam, and scams on the platform. You’ll work with production data alongside experienced engineers, and the models you build will run on live traffic, not just in a notebook. ⚡



If you enjoy working with data, training models, and seeing your work make real decisions for millions of users, this is a great opportunity to kick-start your ML journey! 🚀



🛠️ What You’ll Do



🖼️ Work on image & video ML - fake-profile detection, and spoofed-camera detection on 1:1 video calls.
🎙️ Work on audio ML - abuse and unsafe-speech classification, multilingual and code-mixed ASR for audio rooms.
💬 Work on text ML - harassment, spam, and scam detection across chat in Hindi, English, and regional languages.
📊 Build and clean datasets from production signals and moderation reports, and help define labelling guidelines.
🧠 Train and fine-tune models in Python with PyTorch - including vision backbones, audio encoders, and transformer-based text models.
🎯 Evaluate models with a production mindset - precision at fixed recall, per-language performance, and the real cost of false positives.
⚙️ Collaborate with backend engineers to turn models into inference services and monitor latency and throughput in production.
🔍 Monitor model drift, review misclassifications with the moderation team, and improve models through better data and retraining.
🤝 Explore LLMs for labelling, policy classifiers, and model evaluation.
💻 Write clean, reproducible code and maintain well-documented experiments.


🔎 What We’re Looking For



🏢 Available for a 6-month, in-office internship.
🎓 Strong academic background from a Tier 1 institution, preferably IITs, BITS, NITs, or equivalent institutions.
⭐ High performers may be considered for a PPO based on performance.
🎓 Pursuing a degree in Computer Science, Engineering, Mathematics, Statistics, or a related field.
🐍 Strong Python skills and solid programming fundamentals.
🔥 Working knowledge of PyTorch or TensorFlow - you’ve trained models yourself.
📚 Strong understanding of ML fundamentals - overfitting, regularisation, train/validation/test hygiene, precision vs. recall.
📈 Comfortable with NumPy, pandas, and notebooks for data work.
👁️ Projects in computer vision - CNNs, ViTs, classification/detection - are a plus.
🎧 Exposure to audio ML - spectrograms, ASR, Whisper, or wav2vec-style models - is a plus.
💬 Experience with NLP, transformers, or Hugging Face is a plus.
🛠️ Basics of SQL, Git, or Docker are a plus.
🛡️ Interest in Trust & Safety, content moderation, or responsible AI is a plus.
🚀 Curiosity, ownership, and a strong willingness to learn.




The FRND team operates 6 days a week, with the 1st and 3rd Saturday being working Saturdays.



About FRND



FRND is redefining the way people connect by building social products that are engaging, safe, inclusive, and fun.



We’re a rapidly growing startup building for millions of users and solving unique challenges across social connection and entertainment. Our ambition is bold, and we're looking for people who want to build, experiment, and solve problems at scale.



Why FRND?



🌍 Impact at Scale: Work on products and initiatives that impact millions of users across India and international markets.

🚀 High Ownership: Take on meaningful problems and have the freedom to drive them from idea to execution.

🧠 Learn with the Best: Work closely with founders, leaders, and high-performing teams while solving real business challenges.

💼 Rewarding Journey: Competitive compensation, equity opportunities, and growth that matches your impact.

🎉 Work Hard, Have Fun: We're serious about building great products, but we also believe in enjoying the journey along the way.

💡 Solve Interesting Problems: Work in an environment where curiosity, experimentation, and first-principles thinking are encouraged.
"""
# Defining Jod Description Schema
class JD(BaseModel):
    role:str
    required_skills:list[str]
    preferred_skills:list[str]
    minimum_experience:float| None
    educational_requirements:list[str]
    responsibilities:list[str]

jobd_schema=JD.model_json_schema()

#System prompt to act as HR Assistant follwing the job description schema
system_prompt=f"""
Your a expert HR assistant, your job is to analyze job description and exract stturcture information from them.
Return ONLY valid JSON matching schema:{jobd_schema}
Do not return schema itself
Do not return fields like properties, titles or types
Fill schema with actual information extracted form job description

if min experience is not mentioned return null
if information for a list is missing return empty list
Do not invent information
Must return valid JSON
"""
# User prompt along with requirement of JSON format output
user_prompt=f"""
Analyse the following job description
{job_description}
Must return valid JSON
"""
message_system={
    "role":"system",
    "content":system_prompt
}
message_user={
    "role":"user",
    "content":user_prompt
}
response_format={
    "type":"json_object"
}
messages=[message_system,message_user]
# Getting Response from LLM
response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)
answer=response.choices[0].message.content

# Raw JSON
raw_json=answer
# print(raw_json)

# getting output into json
import json
job_data=json.loads(raw_json)
job=JD(**job_data)
# print(job.role)
# print(job.educational_requirements)
# print(job.minimum_experience)

# Parsing the  resume
# 1. Getting maching schema
class matchresult(BaseModel):
    score:float
    detail:dict
# 2. Getting experience Schema
class experience(BaseModel):
    company:str | None = None
    role:str | None = None
    duration:str | None = None
    description:str | None = None
    skills_used:list[str] = []

# 3. Getting resume schema
class resume(BaseModel):
    name:str | None=None
    email:str | None=None
    phone:str | None=None
    total_experience:float | None=None
    skills:list[str]=[]
    experiences:list[experience]=[]
    education:list[str]=[]
    project:list[str]=[]
resume_schema=resume.model_json_schema()

# Function to calculate the score of resume against job description
def final_score(job,resume):
    match_schema=matchresult.model_json_schema()
    prompt=f"""Youre a HR recruiter
    Compare Candiate's resume with the job description .
    Job Description:
    {job.model_dump_json(indent=2)}
    Candidate Resume:
    {resume.model_dump_json(indent=2)}
    Return the matchings in this schema {match_schema}
    Give me :
    1. Candiate name
    2. Matching skills
    3. Missing Imoprtant skills
    4. Whether experience requirements is met 
    5. Overall match percentage from 0 to 100%
    6. A short final verdict
    Keep the response concise and easy to read 
    Must return valid JSON
    """
    message={
        "role":"user",
        "content":prompt
    }
    messages=[message]
    response_format={
        "type":"json_object"
    }
    response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)
    data=json.loads(response.choices[0].message.content)
    return matchresult(**data)
#  Parsing the resume into Schema from raw text
def parse_resume(resume_text):
    system_prompt=f"""
    You are an expert resume parser
    Extract information from resume based on the meaning, not only based on exact section heading.
    Different resume may have different headings
    for example:
    -Experinece
    -Work Experience
    -Work History
    -Employement
    -Internships
    These all may  contains relevant experience
    Skills may also appears in skill section, work experience, internships or projects
    Return only valid JSON matching this schema
    {resume_schema}

    Important Rules:
    1. Do not invent information
    2. If a value  is not available  return null
    3. if a list has no information return an empty list
    4. Include  internships inside experience
    5. Extract skills mentioned across entire resume
    Must return valid JSON
    """
    user_prompt=f"""
    Parse the following resume:
    {resume_text}
    Must return valid JSON  
    """
    system_message={
        "role":"system",
        "content":system_prompt
    }
    user_message={
        "role":"user",
        "content":user_prompt
    }
    messages=[system_message,user_message]
    response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)
    raw_output=response.choices[0].message.content
    data=json.loads(raw_output)
    return resume(**data)

# Function to read DOCX format resume into resume text(raw)
def read_docx(file_path):
    document=Document(file_path)
    text=""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text+=paragraph.text + "\n"
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text+=cell.text + "\n"
    return text

# Getting resume text(raw) from PDF
def read_pdf(file_path):
    reader=PdfReader(file_path)
    text=""
    for page in reader.pages:
        page_text=page.extract_text()
        if page_text:
            text+=page_text+"\n"
    return text

# choosing between pdf/docs function based on input
def read_resume(file_path):
    if file_path.suffix.lower()==".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower()==".docx":
        return read_docx(file_path)
    else:
        return None

# Final Execution
# getting resume folder Path
resume_folder=Path("resume")
all_results=[] # for storing results
for file_path in resume_folder.iterdir(): #iterating all files in the directory "resume" folder
    if file_path.suffix.lower() not in [".pdf",".docx"]:
        continue                                            # Skipping if resume have other doc type then PDF/DOC
    print("\nProcessing:",file_path.name)
    resume_text=read_resume(file_path)  #Choosing func based on doc type
    parse_resume_result=parse_resume(resume_text) # LLM call for parsing the resume text(raw) into JSON
    time.sleep(5) # Break to prevent API request rate 
    result=final_score(job,parse_resume_result) # Comparing parsed resume text (JSON) with JOB Description
    time.sleep(5) # Break to prevent API request rate
    print("Score:",result.score) #Priting final score individually
    all_results.append({     #Storing all resume results
            "name":parse_resume_result.name,
            "score":result.score,
            "detail":result.detail
    })
all_results.sort(key=lambda candidate:candidate["score"], reverse=True) # Sorting based on score
top_2=all_results[:2] # Getting top 2 candidates
bottom_2=all_results[-2:] # Getting bottom 2 Candidate

print("Top 2 Candidate") # Printing Results of Top 2
for candidate in top_2:
    print(
        candidate["name"],
        "-",
        candidate["score"],"%"
    )
    print(
        candidate["detail"]
    )
print("Bottom 2 Candidate") # Printing Results of Top 2
for candidate in bottom_2:
    print(
        candidate["name"],
        "-",
        candidate["score"],"%"
    )
    print(
        candidate["detail"]
    )
